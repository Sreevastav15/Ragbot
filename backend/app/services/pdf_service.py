# app/services/pdf_service.py
"""
Document text extraction.
Supports: PDF (.pdf), Word (.docx), plain text (.txt)
"""

from __future__ import annotations
import os
from langchain_community.document_loaders import PyPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document


def extract_text(file_path: str) -> list[Document]:
    ext = os.path.splitext(file_path)[-1].lower()

    if ext == ".pdf":
        return _extract_pdf(file_path)
    elif ext == ".docx":
        return _extract_docx(file_path)
    elif ext == ".txt":
        return _extract_txt(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def _extract_pdf(file_path: str) -> list[Document]:
    loader = PyPDFLoader(file_path)
    pages = loader.load()
    chunk_size = min(1500, max(500, len(pages) * 50))
    splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=200)

    chunks = []
    for i, page in enumerate(pages):
        for split in splitter.split_text(page.page_content):
            chunks.append(Document(
                page_content=split,
                metadata={"page_number": i + 1}
            ))
    return chunks


def _extract_docx(file_path: str) -> list[Document]:
    try:
        import docx2txt
        text = docx2txt.process(file_path)
    except ImportError:
        from docx import Document as DocxDocument
        doc = DocxDocument(file_path)
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())

    splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    splits = splitter.split_text(text)
    return [Document(page_content=s, metadata={"page_number": i + 1}) for i, s in enumerate(splits)]


def _extract_txt(file_path: str) -> list[Document]:
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        text = f.read()
    splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    splits = splitter.split_text(text)
    return [Document(page_content=s, metadata={"page_number": i + 1}) for i, s in enumerate(splits)]